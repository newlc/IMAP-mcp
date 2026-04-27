"""Optional text extraction for common attachment formats.

Pure-Python extractors for PDF (pypdf) and Word ``.docx`` (python-docx).
Both libraries are listed under the ``attachments`` extras_require so a
minimal install doesn't pull them in. When unavailable, the extractor
returns ``""`` -- callers should treat the attachment as unindexable, not
fail.

The output text is what goes into the FTS5 ``body`` column for the
attached file. We deliberately keep the formatting minimal (one space
between tokens, no page markers) -- BM25 doesn't care, and short text is
cheaper to store and search.
"""

from __future__ import annotations

import io
import logging
import os
import zipfile
from typing import Optional

logger = logging.getLogger(__name__)

# How many characters of extracted text to keep per attachment. Bounded so
# a malicious or accidental 500-page PDF doesn't blow up the FTS index.
MAX_EXTRACTED_CHARS = 200_000

# Content-Type prefixes / suffixes we know how to handle.
_PDF_TYPES = ("application/pdf", "application/x-pdf")
_DOCX_TYPES = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
)
_PLAIN_TEXT_PREFIXES = ("text/",)


def can_extract(content_type: Optional[str], filename: Optional[str] = None) -> bool:
    """Quick predicate: do we have an extractor for this attachment?"""
    ct = (content_type or "").lower()
    fn = (filename or "").lower()
    if ct in _PDF_TYPES or fn.endswith(".pdf"):
        return True
    if ct in _DOCX_TYPES or fn.endswith(".docx"):
        return True
    if any(ct.startswith(p) for p in _PLAIN_TEXT_PREFIXES):
        return True
    return False


def extract_text(
    data: bytes,
    content_type: Optional[str] = None,
    filename: Optional[str] = None,
) -> str:
    """Best-effort text extraction. Returns ``""`` on unsupported formats
    or when the optional dependency is missing.

    Bounded by :data:`MAX_EXTRACTED_CHARS` to keep the FTS index sane.
    """
    if not data:
        return ""

    ct = (content_type or "").lower()
    fn = (filename or "").lower()

    try:
        if ct in _PDF_TYPES or fn.endswith(".pdf"):
            return _extract_pdf(data)
        if ct in _DOCX_TYPES or fn.endswith(".docx"):
            return _extract_docx(data)
        if any(ct.startswith(p) for p in _PLAIN_TEXT_PREFIXES):
            charset = "utf-8"
            return _truncate(data.decode(charset, errors="replace"))
    except Exception as exc:
        logger.debug("Attachment extraction failed (%s/%s): %s", ct, fn, exc)
    return ""


def _truncate(text: str) -> str:
    if len(text) > MAX_EXTRACTED_CHARS:
        return text[:MAX_EXTRACTED_CHARS]
    return text


def _extract_pdf(data: bytes) -> str:
    try:
        import pypdf
    except ImportError:
        logger.debug(
            "pypdf is not installed; install with `pip install imap-mcp[attachments]` "
            "to extract text from PDF attachments."
        )
        return ""

    reader = pypdf.PdfReader(io.BytesIO(data))
    chunks: list[str] = []
    total_chars = 0
    for page in reader.pages:
        try:
            page_text = page.extract_text() or ""
        except Exception:
            continue
        if not page_text:
            continue
        chunks.append(page_text)
        total_chars += len(page_text)
        if total_chars >= MAX_EXTRACTED_CHARS:
            break
    return _truncate("\n".join(chunks))


def _extract_docx(data: bytes) -> str:
    try:
        import docx  # python-docx
    except ImportError:
        logger.debug(
            "python-docx is not installed; install with "
            "`pip install imap-mcp[attachments]` to extract text from .docx attachments."
        )
        return ""

    # python-docx accepts a file-like object. Ensure the bytes really are a
    # docx (zip with word/document.xml).
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            if "word/document.xml" not in zf.namelist():
                return ""
    except zipfile.BadZipFile:
        return ""

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception:
        return ""

    paragraphs = [p.text for p in document.paragraphs if p.text]
    # Tables are commonly used for layout in real docs; pull their cell text too.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paragraphs.append(cell.text)
    return _truncate("\n".join(paragraphs))
