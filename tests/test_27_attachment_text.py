"""Tests for attachment text extraction (#5).

PDF and DOCX text gets extracted on store_attachment and appended to the
email's FTS row, so search_emails_fts finds words inside attachments.
"""

from __future__ import annotations

import io
import zipfile

import pytest

from imap_mcp.attachments import (
    MAX_EXTRACTED_CHARS,
    can_extract,
    extract_text,
)
from imap_mcp.cache import EmailCache


# ---------------------------------------------------------------------------
# can_extract / extract_text predicate
# ---------------------------------------------------------------------------


class TestCanExtract:
    @pytest.mark.parametrize("ct,fn,expected", [
        ("application/pdf", None, True),
        ("application/x-pdf", None, True),
        (None, "report.pdf", True),
        ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",
         None, True),
        (None, "memo.docx", True),
        ("text/plain", None, True),
        ("text/html", None, True),
        ("image/png", None, False),
        ("application/octet-stream", None, False),
        ("application/zip", "x.zip", False),
        (None, None, False),
    ])
    def test_predicate(self, ct, fn, expected):
        assert can_extract(ct, fn) is expected

    def test_extract_empty_returns_empty(self):
        assert extract_text(b"", "text/plain") == ""
        assert extract_text(None, "text/plain") == ""


class TestExtractPlainText:
    def test_basic(self):
        text = extract_text(b"hello world", "text/plain")
        assert text == "hello world"

    def test_truncated_above_max(self):
        big = ("x" * (MAX_EXTRACTED_CHARS + 1000)).encode()
        out = extract_text(big, "text/plain")
        assert len(out) <= MAX_EXTRACTED_CHARS

    def test_unsupported_returns_empty(self):
        assert extract_text(b"\x89PNG", "image/png") == ""

    def test_unknown_filename_extension(self):
        assert extract_text(b"x", None, "file.unknown") == ""


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def _make_docx(paragraphs: list[str]) -> bytes:
    """Build a minimal valid .docx in-memory."""
    import docx
    doc = docx.Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestExtractDocx:
    def test_paragraphs_extracted(self):
        data = _make_docx(["First paragraph", "Second paragraph"])
        text = extract_text(
            data,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "memo.docx",
        )
        assert "First paragraph" in text
        assert "Second paragraph" in text

    def test_table_cells_extracted(self):
        import docx
        doc = docx.Document()
        doc.add_paragraph("intro")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "left-cell"
        table.cell(0, 1).text = "right-cell"
        buf = io.BytesIO()
        doc.save(buf)
        text = extract_text(buf.getvalue(), filename="t.docx")
        assert "left-cell" in text
        assert "right-cell" in text

    def test_corrupt_zip_returns_empty(self):
        # Looks like a docx by extension but isn't a valid zip.
        out = extract_text(b"not a zip", filename="x.docx")
        assert out == ""

    def test_zip_without_word_xml_returns_empty(self):
        # A valid zip but missing word/document.xml.
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as zf:
            zf.writestr("hello.txt", "not a docx")
        out = extract_text(buf.getvalue(), filename="x.docx")
        assert out == ""


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def _make_pdf(text: str) -> bytes:
    """Build a minimal PDF via pypdf (Page + Tj operator)."""
    import pypdf
    writer = pypdf.PdfWriter()
    writer.add_blank_page(width=612, height=792)
    # Inject a simple content stream so the page has extractable text.
    page = writer.pages[0]
    page.merge_page  # ensures page is initialized
    # Direct text injection via PageObject.add_transformation isn't
    # available, so fabricate a content stream manually.
    from pypdf.generic import (
        ArrayObject, ContentStream, DecodedStreamObject,
        DictionaryObject, FloatObject, NameObject, NumberObject,
    )
    cs_text = (
        f"BT /F1 24 Tf 100 700 Td ({text}) Tj ET"
    ).encode()
    stream = DecodedStreamObject()
    stream.set_data(cs_text)
    page[NameObject("/Contents")] = stream
    # Add a font so /F1 resolves.
    font = DictionaryObject()
    font[NameObject("/Type")] = NameObject("/Font")
    font[NameObject("/Subtype")] = NameObject("/Type1")
    font[NameObject("/BaseFont")] = NameObject("/Helvetica")
    fonts = DictionaryObject()
    fonts[NameObject("/F1")] = font
    resources = DictionaryObject()
    resources[NameObject("/Font")] = fonts
    page[NameObject("/Resources")] = resources

    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


class TestExtractPdf:
    def test_extracts_text(self):
        pdf = _make_pdf("Hello pdf world")
        text = extract_text(pdf, "application/pdf", "x.pdf")
        # pypdf's text extraction may add whitespace or split words; the key
        # tokens should appear.
        assert "Hello" in text or "pdf" in text or "world" in text

    def test_corrupt_pdf_returns_empty(self):
        out = extract_text(b"%PDF-not-really", "application/pdf", "x.pdf")
        assert out == ""


# ---------------------------------------------------------------------------
# Cache integration: store_attachment populates FTS, search finds it
# ---------------------------------------------------------------------------


class TestCacheAttachmentFts:
    @pytest.fixture
    def cache(self, tmp_cache_db):
        return EmailCache(tmp_cache_db, encrypted=False, keyring_username="t-attfts")

    def _store_email(self, cache, uid: int, subject: str, body: str = ""):
        cache.store_email(
            "INBOX", uid,
            {
                "message_id": f"<{uid}@x>", "subject": subject,
                "from_address": {"email": "a@x.com"}, "to_addresses": [],
                "cc_addresses": [], "date": None, "flags": [], "size": 0,
            },
            {"text": body, "html": None},
        )

    def test_plain_text_attachment_searchable(self, cache):
        self._store_email(cache, 1, subject="Quarterly review")
        cache.store_attachment(
            "INBOX", 1, 0, "notes.txt", "text/plain", 50,
            b"unique-token-xyzzy lives inside this attachment",
        )
        rows = cache.fts_search("xyzzy")
        assert len(rows) == 1
        assert rows[0]["uid"] == 1

    def test_docx_attachment_searchable(self, cache):
        self._store_email(cache, 2, subject="Spec review")
        docx_bytes = _make_docx(["Project alpha milestone deadlines"])
        cache.store_attachment(
            "INBOX", 2, 0, "spec.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            len(docx_bytes), docx_bytes,
        )
        rows = cache.fts_search("milestone")
        assert any(r["uid"] == 2 for r in rows)

    def test_unsupported_attachment_does_not_break(self, cache):
        self._store_email(cache, 3, subject="Image attached")
        cache.store_attachment(
            "INBOX", 3, 0, "photo.jpg", "image/jpeg", 50, b"\xFF\xD8\xFF\xE0",
        )
        # Searching for the binary doesn't match, but the email itself is
        # still indexed by subject.
        assert any(r["uid"] == 3 for r in cache.fts_search("Image"))

    def test_subject_outranks_attachment_via_bm25(self, cache):
        # UID 1: word in subject only.
        self._store_email(cache, 1, subject="invoice tracker")
        cache.store_attachment("INBOX", 1, 0, "n.txt", "text/plain", 5, b"")
        # UID 2: word in attachment only.
        self._store_email(cache, 2, subject="random subject")
        cache.store_attachment(
            "INBOX", 2, 0, "n.txt", "text/plain", 50,
            b"this attachment mentions invoice once",
        )
        rows = cache.fts_search("invoice")
        # bm25 weights subject=5, attachments=1.2 -- subject wins.
        assert rows[0]["uid"] == 1


class TestFtsSchemaMigration:
    def test_old_six_column_fts_gets_rebuilt(self, tmp_cache_db):
        # Simulate an older cache: pre-create FTS with the old 6-column
        # shape. The regular emails / mailbox_meta / etc. tables are
        # created from _SCHEMA on the first EmailCache open.
        import sqlite3
        conn = sqlite3.connect(tmp_cache_db)
        conn.executescript(
            """
            CREATE VIRTUAL TABLE emails_fts USING fts5(
              mailbox UNINDEXED, uid UNINDEXED,
              subject, body, from_address, to_address
            );
            """
        )
        conn.commit()
        conn.close()

        # Opening via EmailCache should detect the 6-column FTS and rebuild.
        cache = EmailCache(tmp_cache_db, encrypted=False, keyring_username="t-mig")
        # Now we should have 7 columns.
        cur = cache.conn.execute("SELECT * FROM emails_fts LIMIT 0")
        assert len(cur.description) == 7
